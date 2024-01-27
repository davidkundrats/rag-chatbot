import PyPDF2
import docx
import os
from langchain.text_splitter import CharacterTextSplitter
from langchain_openai import AzureOpenAIEmbeddings

def get_pdf_text(pdf_list): 
    text = ""
    try:
        for pdf in pdf_list: 
            with open(pdf, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text()
    except PyPDF2.utils.PdfReadError as e:
        print(f"An error occurred while reading the PDF: {e}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
    return text

def get_docx_text(docx_list):
    text = []
    try:
        for document in docx_list:
            doc = docx.Document(document)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(f"{cell.text.strip()} ")
            for para in doc.paragraphs:
                text.append(para.text.strip())
    except docx.exceptions.PackageNotFoundError as e:
        print(f"An error occurred while opening the Word document: {e}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
    return ' '.join(text)

#chunk text
def get_chunks(text): 
    text_splitter = CharacterTextSplitter(
        separator='\n', 
        chunk_size = 1000, 
        chunk_overlap = 200, 
        length_function= len
    ) 
    chunks = text_splitter.split_text(text)
    return clean_chunks(chunks)
    
#chunk helper func
def clean_text(text):
    text = text.replace('\xa0', ' ')
    text = text.replace('\n', ' ')
    return ' '.join(text.split())
def clean_chunks(chunks): 
    return [clean_text(chunk) for chunk in chunks]
#get embeddings
import os
def get_embeddings(chunk_text): 
    try:
        embedding_list = []
        embeddings = AzureOpenAIEmbeddings(
            azure_deployment=os.getenv("AZURE_DEPLOYMENT"),
            openai_api_version="2023-05-15",
        )
        for chunk in chunk_text: 
            embedding_list.append(embeddings.embed_query(chunk))
    except Exception as e: 
        print(f'An error occured while embedding texts. {e}')
    return embedding_list
#generate metadata 
def generate_metadata_for_vector(index, text_chunk, category="General HR"):
    return {
        "chunk_id": f"HR_Chunk_{index:03d}", 
        "category": category, 
        "text":text_chunk,
        "description": f"Chunk {index} of HR text data"
    }
#pinecone upsert method
from pinecone import Pinecone
def pinecone_upsert(embeddings, chunk_text, index_name="knowledge-base"):
    try:
        pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        index = pc.Index(index_name)
    
        data = []
        for i, embedding in enumerate(embeddings):
            metadata = generate_metadata_for_vector(i, chunk_text[i])
            data.append((f"chunk_{i}", embedding, metadata))

        index.upsert(vectors=data)
    except Exception as e: 
        print(f'An error occured while upserting data: {e}')

mnt = "/dbfs/mnt/data"
docxs = [os.path.join(mnt, file) for file in os.listdir(mnt) if file.endswith('.docx')]
pdfs = [os.path.join(mnt, file) for file in os.listdir(mnt) if file.endswith('.pdf')]

pdf_text = get_pdf_text(pdfs)
docx_text = get_docx_text(docxs)
full_doc_text = pdf_text + docx_text
chunk_text = get_chunks(full_doc_text)
#getting embeddings
embeddings_list = get_embeddings(chunk_text)
#push our vector to pinecone!
pinecone_upsert(embeddings_list, chunk_text)
